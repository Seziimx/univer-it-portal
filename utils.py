from docx import Document
from io import BytesIO

def generate_word_report(zayavki):
    doc = Document()
    doc.add_heading('Отчёт по заявкам', 0)

    # Заголовок таблицы
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Сотрудник'
    hdr_cells[1].text = 'Тип'
    hdr_cells[2].text = 'Описание'
    hdr_cells[3].text = 'Дата'

    # Заполнение таблицы
    for z in zayavki:
        row_cells = table.add_row().cells
        row_cells[0].text = z.user.username
        row_cells[1].text = z.type
        row_cells[2].text = z.description
        row_cells[3].text = z.created_at.strftime('%d.%m.%Y %H:%M')

    # Сохранение в BytesIO
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from io import BytesIO

def generate_pdf_report(zayavki):
    byte_io = BytesIO()
    c = canvas.Canvas(byte_io, pagesize=letter)
    c.setFont("Helvetica", 10)

    # Заголовок
    c.drawString(100, 750, 'Отчёт по заявкам')

    # Колонки
    y_position = 730
    c.drawString(100, y_position, 'Сотрудник')
    c.drawString(200, y_position, 'Тип')
    c.drawString(300, y_position, 'Описание')
    c.drawString(400, y_position, 'Дата')

    # Заполнение таблицы
    y_position -= 20
    for z in zayavki:
        c.drawString(100, y_position, z.user.username)
        c.drawString(200, y_position, z.type)
        c.drawString(300, y_position, z.description)
        c.drawString(400, y_position, z.created_at.strftime('%d.%m.%Y %H:%M'))
        y_position -= 20

    c.showPage()
    c.save()

    byte_io.seek(0)
    return byte_io
